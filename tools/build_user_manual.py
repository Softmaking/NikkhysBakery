from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "manuales"
DOCX_PATH = OUT_DIR / "manual-rapido-usuario-nikkhysbakery.docx"
LOGO_PATH = ROOT / "NikkhysBakery-Front" / "public" / "assets" / "brand" / "logo-public-header.png"


BRAND = RGBColor(126, 77, 78)
BRAND_DARK = RGBColor(62, 45, 43)
MUTED = RGBColor(116, 99, 96)
LIGHT = "F7F1EC"
LIGHT_2 = "FBF8F5"
TABLE_HEADER = "EDE2DC"
GREEN_FILL = "EAF4EE"
WARN_FILL = "FFF4E6"
RED_FILL = "FCEEEE"
BORDER = "E3D5CE"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, val in (("top", "90"), ("bottom", "90"), ("start", "130"), ("end", "130")):
                child = margins.find(qn(f"w:{side}"))
                if child is None:
                    child = OxmlElement(f"w:{side}")
                    margins.append(child)
                child.set(qn("w:w"), val)
                child.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)


def set_run(run, size=10.5, color=BRAND_DARK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", style=None, size=10.5, color=BRAND_DARK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    color = BRAND if level < 3 else MUTED
    size = {1: 15.5, 2: 13, 3: 11.5}.get(level, 11)
    set_run(run, size=size, color=color, bold=True)
    return p


def add_callout(doc, title, body, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run(r, size=10.5, color=BRAND, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_run(r2, size=10, color=BRAND_DARK)
    add_para(doc, "", after=4)


def add_data_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], TABLE_HEADER)
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run(r, size=9.2, color=BRAND, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_borders(cells[idx])
            set_cell_shading(cells[idx], LIGHT_2)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(value)
            set_run(r, size=9.4, color=BRAND_DARK)
    add_para(doc, "", after=3)
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.36)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BRAND_DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, before, after in (
        ("Heading 1", 15.5, 15, 7),
        ("Heading 2", 13, 11, 5),
        ("Heading 3", 11.5, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BRAND
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Manual rapido de uso NikkhysBakery")
    set_run(r, size=8.5, color=MUTED)


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    if LOGO_PATH.exists():
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(0.62))

    add_para(doc, "NikkhysBakery", size=24, color=BRAND, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Manual rapido de uso", size=17, color=BRAND_DARK, bold=True, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Stock, recetas, produccion y ventas", size=12.5, color=MUTED, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(
        doc,
        "Objetivo",
        "Esta guia explica el flujo operativo diario en pocas paginas: revisar stock, mantener recetas, producir y vender sin perder trazabilidad.",
        fill=LIGHT,
    )
    add_data_table(
        doc,
        ["Tema", "Que resuelve"],
        [
            ("Stock", "Saber que materia prima y producto vendible hay por sucursal."),
            ("Recetas", "Definir que insumos consume cada preparacion y que resultado genera."),
            ("Produccion", "Transformar insumos en producto disponible o stock vendible."),
            ("Ventas", "Vender en POS respetando stock, disponibilidad y cierre de caja."),
        ],
        [1900, 7460],
    )
    add_para(doc, f"Version operativa: mayo 2026", size=9.4, color=MUTED, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def build_manual():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    cover(doc)

    add_heading(doc, "1. Antes de operar: idea clave", 1)
    add_para(
        doc,
        "El sistema separa inventario de materia prima, recetas, produccion y productos de venta. "
        "Esa separacion evita confundir harina, leche o cafe con los productos que el cliente compra en el POS.",
    )
    add_data_table(
        doc,
        ["Concepto", "Uso correcto", "Ejemplo"],
        [
            ("Materia prima", "Insumos que se compran, reciben y consumen.", "Harina, mantequilla, cafe molido, leche."),
            ("Producto de venta", "Producto visible en POS para vender al cliente.", "Pan amasado, capuchino, cheesecake."),
            ("Receta", "Define insumos, rendimiento teorico y producto asociado.", "Pan amasado consume harina, agua, sal y levadura."),
            ("Produccion", "Consume materia prima y habilita stock o disponibilidad.", "Completar produccion de pan para venderlo."),
        ],
        [1700, 4750, 2910],
    )
    add_callout(
        doc,
        "Regla practica",
        "Si quieres producir, revisa receta e inventario. Si quieres vender, revisa producto de venta y disponibilidad por sucursal.",
        fill=GREEN_FILL,
    )

    add_heading(doc, "2. Rutina diaria recomendada", 1)
    add_data_table(
        doc,
        ["Momento", "Accion", "Resultado esperado"],
        [
            ("Inicio", "Elegir sucursal activa y revisar stock bajo o productos agotados.", "El equipo sabe que puede producir y vender."),
            ("Antes de producir", "Planificar produccion desde una receta activa.", "Se ven insumos requeridos, disponibles y faltantes."),
            ("Durante el dia", "Vender en POS y marcar agotado cuando corresponda.", "No se venden productos sin stock o sin disponibilidad."),
            ("Cierre", "Revisar ventas, movimientos y producciones completadas.", "Queda trazabilidad para reportes y reposicion."),
        ],
        [1500, 5000, 2860],
    )

    add_heading(doc, "3. Stock e inventario", 1)
    add_heading(doc, "Que revisar", 2)
    add_para(
        doc,
        "En inventario se administran materias primas, ubicaciones, ingresos, ajustes, mermas y stock por sucursal. "
        "Para operacion diaria, lo mas importante es confirmar que la sucursal tenga stock suficiente antes de producir.",
    )
    add_data_table(
        doc,
        ["Operacion", "Cuando usarla", "Cuidado"],
        [
            ("Ingreso", "Cuando entra compra o reposicion de insumos.", "Registrar costo total y unidad correcta."),
            ("Ajuste", "Cuando el conteo real no coincide con el sistema.", "Usarlo con motivo claro; no reemplaza una venta."),
            ("Merma", "Cuando se pierde producto o insumo.", "Registrar para no ocultar diferencias de stock."),
            ("Movimientos", "Cuando necesitas investigar cambios.", "Filtrar por sucursal, fecha e item."),
        ],
        [1600, 4300, 3460],
    )
    add_callout(
        doc,
        "Sucursal",
        "El stock operativo se trabaja por sucursal. Un administrador puede ver stock global, pero las entradas, mermas y producciones deben quedar asociadas a una sucursal concreta.",
        fill=LIGHT,
    )

    add_heading(doc, "4. Recetas", 1)
    add_para(
        doc,
        "Una receta sirve para saber que insumos se consumen, cuanto cuesta producir y que producto queda disponible para vender. "
        "Las recetas se trabajan por versiones: una version puede estar en borrador y solo una version queda activa para operar.",
    )
    add_data_table(
        doc,
        ["Paso", "Que hacer", "Verificacion"],
        [
            ("Crear receta", "Indicar nombre, tipo y producto asociado para venta cuando corresponda.", "El producto asociado debe existir en Productos."),
            ("Agregar insumos", "Registrar materias primas y cantidades de la preparacion.", "Las unidades deben coincidir o tener conversion."),
            ("Definir rendimiento", "Indicar cuanto produce la receta para costeo y planificacion.", "Ej.: 30 unidades teoricas o una preparacion."),
            ("Crear version", "Guardar la version de receta.", "Queda como borrador hasta activarla."),
            ("Activar version", "Activar cuando este completa y validada.", "Produccion y costos usaran esa version."),
        ],
        [1300, 4700, 3360],
    )
    add_callout(
        doc,
        "Producto por disponibilidad",
        "Para productos como pan vendido por monto, el rendimiento se usa para costeo. La venta no descuenta unidades; el equipo marca el producto como agotado cuando ya no queda.",
        fill=WARN_FILL,
    )

    add_heading(doc, "5. Produccion", 1)
    add_para(
        doc,
        "Produccion transforma materia prima en producto terminado o en disponibilidad para vender. Crear una orden no consume stock; el consumo ocurre al completar la produccion.",
    )
    add_data_table(
        doc,
        ["Accion", "Que revisar", "Que pasa en sistema"],
        [
            ("Planificar", "Receta activa, sucursal y cantidad a producir.", "Muestra requeridos, disponibles y faltantes."),
            ("Crear orden", "Notas, sucursal y cantidad planificada.", "La orden queda registrada para seguimiento."),
            ("Confirmar", "Que se ejecutara la produccion.", "La orden avanza de estado."),
            ("Completar", "Que se produjo realmente.", "Consume insumos y aumenta stock vendible o disponibilidad."),
            ("Cancelar", "Solo si la produccion no seguira.", "La orden queda cerrada sin completar consumo."),
        ],
        [1550, 4300, 3510],
    )
    add_callout(
        doc,
        "Si no deja producir",
        "Revisa que exista sucursal activa, receta con version activa, insumos configurados, producto asociado y stock suficiente de materias primas.",
        fill=RED_FILL,
    )

    add_heading(doc, "6. Ventas / POS", 1)
    add_para(
        doc,
        "El POS se usa para crear ordenes, agregar productos, enviar a cocina cuando corresponda y cerrar la venta. "
        "El sistema valida el stock o disponibilidad antes de vender y guarda el costo historico al cerrar.",
    )
    add_data_table(
        doc,
        ["Tipo de producto", "Como se vende", "Efecto en stock"],
        [
            ("Stock vendible", "Se agrega al POS con cantidad.", "Descuenta producto terminado al cerrar la orden."),
            ("Disponibilidad", "Se vende mientras este disponible.", "No descuenta unidades; se marca agotado manualmente."),
            ("Receta directa", "Se prepara al momento.", "Consume insumos de la receta activa al cerrar."),
            ("Sin inventario", "Se vende sin control de stock.", "No mueve inventario."),
        ],
        [2100, 3650, 3610],
    )
    add_callout(
        doc,
        "Cierre de venta",
        "El inventario se descuenta al cerrar la orden, no solo al agregar productos. Si el producto no tiene stock o esta agotado, no debe venderse.",
        fill=GREEN_FILL,
    )

    add_heading(doc, "7. Flujo completo recomendado", 1)
    add_data_table(
        doc,
        ["Orden", "Modulo", "Accion breve"],
        [
            ("1", "Inventario", "Registrar ingresos de materia prima con costo y sucursal correctos."),
            ("2", "Productos", "Configurar productos de venta y su modo de inventario."),
            ("3", "Recetas", "Crear receta, asociar producto, agregar insumos y activar version."),
            ("4", "Produccion", "Planificar, crear, confirmar y completar la orden."),
            ("5", "POS", "Vender productos disponibles y cerrar ordenes."),
            ("6", "Reportes", "Revisar ventas, stock, movimientos y produccion para corregir desviaciones."),
        ],
        [1000, 1900, 6460],
    )

    add_heading(doc, "8. Problemas comunes", 1)
    add_data_table(
        doc,
        ["Situacion", "Que revisar primero"],
        [
            ("No aparece una receta en produccion", "Debe tener version activa y producto asociado/resultado valido."),
            ("No deja activar una receta", "Faltan insumos, rendimiento o datos requeridos de la version."),
            ("No deja vender un producto", "Puede estar sin stock vendible, agotado, inactivo o sin sucursal correcta."),
            ("El stock no coincide", "Revisar movimientos, ubicacion, sucursal, ajustes y mermas."),
            ("Produccion no descuenta insumos", "Verificar que la orden se completo; crear la orden no consume stock."),
        ],
        [3600, 5760],
    )

    add_heading(doc, "9. Reglas de oro", 1)
    add_data_table(
        doc,
        ["Regla", "Aplicacion diaria"],
        [
            ("No mezclar inventario y POS", "Inventario es materia prima; POS vende productos."),
            ("Siempre operar con sucursal correcta", "Evita stock duplicado, ventas bloqueadas y reportes incorrectos."),
            ("No producir sin receta activa", "La receta activa es la fuente para consumo y costo."),
            ("Registrar mermas y ajustes con motivo", "Permite explicar diferencias reales."),
            ("Cerrar ventas al final del flujo", "El cierre deja trazabilidad y mueve stock cuando corresponde."),
        ],
        [2800, 6560],
    )
    add_para(
        doc,
        "Este manual es una guia rapida. Para dudas de permisos, configuracion avanzada o reportes, consultar al administrador del sistema.",
        size=9.5,
        color=MUTED,
        italic=True,
        before=6,
        after=0,
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_manual()
